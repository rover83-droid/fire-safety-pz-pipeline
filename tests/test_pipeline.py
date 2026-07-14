from __future__ import annotations

import json
import shutil
import unittest
import uuid
from pathlib import Path

from mpb_pz_flow import io
from mpb_pz_flow.assembler import assemble_draft, finalize_markdown
from mpb_pz_flow.codex_native import (
    can_run_agent,
    infer_fkp,
    missing_for_agent,
    read_conversation_state,
    record_passport_answer,
    start_codex_project,
    status_card,
)
from mpb_pz_flow.demo import create_demo
from mpb_pz_flow.models import MatrixEntry, NormEntry, Stage
from mpb_pz_flow.pipeline import audit_binding_issues, audit_project, compute_stage, init_project, refresh_stage
from mpb_pz_flow.standards import FKP_TABLE
from mpb_pz_flow import corpus
from mpb_pz_flow.validators import (
    validate_all,
    validate_draft,
    validate_editions,
    validate_norms,
    validate_passport,
    validate_quotes,
    validate_volume_structure,
)
from mpb_pz_flow.gui_server import list_projects, project_snapshot
from mpb_pz_flow.volume_structure import canonical_section_titles_for_fkp


class PipelineTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def test_demo_project_passes_audit_and_finalizes(self) -> None:
        project = self.temp_project_parent() / "demo"
        create_demo(project)

        state = refresh_stage(project)
        self.assertEqual(Stage.AUDIT_PASSED, state.stage)

        report = io.read_json(io.artifact_dir(project) / "audit_report.json")
        self.assertTrue(report["passed"])

        final = finalize_markdown(project)
        self.assertIn("Наружное ВПС и проезды", final)
        self.assertIn("СП 8.13130.2020, п. 5.2, табл. 2", final)
        self.assertNotIn("{norm:", final)
        self.assertNotIn("[Тип ", final)

    def test_init_project_requires_remaining_passport_fields(self) -> None:
        project = self.temp_project_parent() / "shop"
        init_project(project, "F3.1", "Наружное ВПС и проезды", "Магазин")

        issues = validate_passport(project)
        codes = {issue.code for issue in issues}
        self.assertIn("passport.required", codes)

    def test_draft_rejects_non_applicable_norm(self) -> None:
        project = self.temp_project_parent() / "demo"
        create_demo(project)

        matrix_path = io.artifact_dir(project) / "applicability_matrix.jsonl"
        rows = io.read_jsonl(matrix_path)
        rows[0]["status"] = "неприменимо"
        io.write_jsonl(matrix_path, rows)

        issues = validate_draft(project)
        codes = {issue.code for issue in issues}
        self.assertIn("draft.non_applicable_norm", codes)

    def test_draft_rejects_internal_norm_ref_without_visible_citation(self) -> None:
        project = self.temp_project_parent() / "visible-citation"
        create_demo(project)

        draft_path = io.artifact_dir(project) / "draft.md"
        draft_path.write_text(
            "# Наружное ВПС и проезды\n\n"
            "[Тип А] Наружное противопожарное водоснабжение предусматривается от пожарных гидрантов. {norm:sp8-5-2-tab2}\n",
            encoding="utf-8",
        )

        issues = validate_draft(project)
        codes = {issue.code for issue in issues}
        self.assertIn("draft.visible_norm_reference", codes)

    def test_draft_accepts_visible_document_year_and_point_citation(self) -> None:
        project = self.temp_project_parent() / "visible-citation-ok"
        create_demo(project)

        issues = validate_draft(project)
        codes = {issue.code for issue in issues}
        self.assertNotIn("draft.visible_norm_reference", codes)

    def test_gui_snapshot_reports_demo_state(self) -> None:
        project = Path.cwd() / "projects" / "demo"
        if not (project / "state.json").exists():
            create_demo(project)

        snapshot = project_snapshot(project)
        self.assertEqual("demo", snapshot["name"])
        self.assertIn("state", snapshot)
        self.assertIn("issues", snapshot)

        names = {project["name"] for project in list_projects()}
        self.assertIn("demo", names)

    def test_codex_native_start_infers_fkp_and_asks_one_question(self) -> None:
        project = self.temp_project_parent() / "codex-shop"
        result = start_codex_project(
            project,
            object_name="Магазин у дома",
            fkp=None,
            section="Наружное ВПС и проезды",
            standards_dir="standards",
            description="Небольшой торговый объект",
        )

        self.assertEqual("Магазин у дома", result["status_card"]["Проект"])
        self.assertIn("next_question", result)

        conversation = read_conversation_state(project)
        self.assertIsNotNone(conversation.expected_passport_field)
        self.assertIsNotNone(conversation.last_question)

        answer = record_passport_answer(project, "III")
        self.assertEqual("fire_resistance_degree", answer["saved_field"])
        self.assertIn("Паспорт", answer["status_card"])

    def test_codex_native_agent_gates_report_blockers(self) -> None:
        project = self.temp_project_parent() / "codex-gates"
        start_codex_project(
            project,
            object_name="Магазин",
            fkp="F3.1",
            section="Наружное ВПС и проезды",
        )

        self.assertFalse(can_run_agent(project, 1))
        self.assertTrue(any("standards" in blocker for blocker in missing_for_agent(project, 1)))
        self.assertFalse(can_run_agent(project, 2))
        self.assertFalse(can_run_agent(project, 3))

        card = status_card(project)
        self.assertIn("Следующее действие", card)

    def test_f5_1_volume_structure_is_available_to_agents(self) -> None:
        titles = canonical_section_titles_for_fkp("F5.1")
        self.assertIn("Основные положения", titles)
        self.assertIn("Инженерные системы противопожарной защиты", titles)
        self.assertIn("Система пожарной сигнализации", titles)
        self.assertIn("Система удаления продуктов горения при пожаре", titles)

    def test_f5_1_full_volume_validator_reports_missing_sections(self) -> None:
        project = self.temp_project_parent() / "f5-structure"
        init_project(project, "F5.1", "Том 9. Мероприятия по обеспечению пожарной безопасности", "Производство")
        draft_path = io.artifact_dir(project) / "draft.md"
        draft_path.write_text(
            "# Том 9. Мероприятия по обеспечению пожарной безопасности\n\n"
            "<!-- volume_structure: F5.1 -->\n\n"
            "## Основные положения\n\n"
            "## Инженерные системы противопожарной защиты\n",
            encoding="utf-8",
        )

        issues = validate_volume_structure(project)
        codes = {issue.code for issue in issues}
        self.assertIn("volume_structure.missing_section", codes)

    def test_refresh_stage_regresses_when_artifact_removed(self) -> None:
        project = self.temp_project_parent() / "demo-regression"
        create_demo(project)
        self.assertEqual(Stage.AUDIT_PASSED, refresh_stage(project).stage)

        (io.artifact_dir(project) / "norms.jsonl").unlink()
        self.assertEqual(Stage.PASSPORT_READY, refresh_stage(project).stage)

    def test_audit_history_preserved(self) -> None:
        project = self.temp_project_parent() / "audit-history"
        create_demo(project)
        audit_project(project)

        history = io.artifact_dir(project) / "audit_history"
        self.assertTrue(history.exists())
        self.assertTrue(list(history.glob("audit_report_*.json")))

    def test_archive_previous_audit_survives_corrupted_report(self) -> None:
        project = self.temp_project_parent() / "audit-corrupt"
        create_demo(project)
        (io.artifact_dir(project) / "audit_report.json").write_text("{broken", encoding="utf-8")

        report = audit_project(project)
        self.assertTrue(report["passed"])

    def test_source_file_validation_warns_on_missing(self) -> None:
        project = self.temp_project_parent() / "missing-source"
        create_demo(project)
        shutil.rmtree(project / "standards")

        issues = validate_norms(project)
        warnings = [issue for issue in issues if issue.code == "norms.source_file_missing"]
        self.assertTrue(warnings)
        self.assertTrue(all(issue.severity == "warning" for issue in warnings))

    def test_fkp_f1_f2_f4_supported(self) -> None:
        expected = {
            "F1.1", "F1.2", "F1.3", "F1.4",
            "F2.1", "F2.2", "F2.3", "F2.4",
            "F3.1", "F3.2", "F3.3", "F3.4", "F3.5", "F3.6",
            "F4.1", "F4.2", "F4.3", "F4.4",
            "F5.1", "F5.2", "F5.3",
        }
        self.assertEqual(expected, set(FKP_TABLE))

    def test_infer_fkp_new_classes(self) -> None:
        self.assertEqual("F1.1", infer_fkp("детский сад"))
        self.assertEqual("F1.2", infer_fkp("гостиница"))
        self.assertEqual("F2.1", infer_fkp("кинотеатр"))
        self.assertEqual("F4.1", infer_fkp("школа"))
        self.assertEqual("F4.4", infer_fkp("офис"))
        self.assertEqual("F5.3", infer_fkp("подземный паркинг"))

    def test_init_f4_project(self) -> None:
        project = self.temp_project_parent() / "f4-project"
        state = init_project(project, "F4.2", "Общие положения", "Учебный корпус")
        self.assertEqual("F4.2", state.fkp)

    def test_init_project_refuses_to_overwrite_existing(self) -> None:
        project = self.temp_project_parent() / "overwrite"
        init_project(project, "F3.1", "Общие положения", "Магазин")
        with self.assertRaises(FileExistsError):
            init_project(project, "F3.1", "Общие положения", "Магазин")

    def test_init_project_force_reinitializes(self) -> None:
        project = self.temp_project_parent() / "force"
        init_project(project, "F3.1", "Общие положения", "Магазин")
        marker = project / "marker.txt"
        marker.write_text("old", encoding="utf-8")

        init_project(project, "F4.2", "Общие положения", "Учебный корпус", force=True)
        self.assertFalse(marker.exists())
        self.assertEqual("F4.2", io.read_state(project).fkp)

    def test_validate_all_does_not_skip_on_warnings(self) -> None:
        project = self.temp_project_parent() / "warnings"
        create_demo(project)
        shutil.rmtree(project / "standards")
        issues = validate_all(project)
        codes = {issue.code for issue in issues}
        self.assertIn("norms.source_file_missing", codes)
        self.assertNotIn("matrix.missing", codes)

    def test_atomic_write_does_not_leave_temp_files(self) -> None:
        project = self.temp_project_parent() / "atomic"
        io.write_json(project / "state.json", {"ok": True})
        self.assertFalse(list(project.glob("*.tmp")))


class Phase0IntegrityTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def test_handwritten_audit_report_is_not_trusted(self) -> None:
        project = self.temp_project_parent() / "fabricated-audit"
        create_demo(project)

        fabricated = {
            "passed": True,
            "verdict": "КОНВЕЙЕР ЗАВЕРШЁН",
            "level_1": [],
            "level_2": [],
            "generated_at": "2026-01-01T00:00:00+00:00",
        }
        audit_path = io.artifact_paths(project).audit_report
        audit_path.write_text(json.dumps(fabricated, ensure_ascii=False), encoding="utf-8")

        problems = audit_binding_issues(project)
        self.assertTrue(problems)
        self.assertEqual(Stage.DRAFT_READY, refresh_stage(project).stage)
        with self.assertRaises(RuntimeError):
            finalize_markdown(project)

    def test_audit_invalidated_when_draft_changes_after_audit(self) -> None:
        project = self.temp_project_parent() / "stale-audit"
        create_demo(project)
        self.assertEqual(Stage.AUDIT_PASSED, refresh_stage(project).stage)

        draft_path = io.artifact_paths(project).draft
        draft_path.write_text(
            draft_path.read_text(encoding="utf-8")
            + "\n[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 принято дополнительное проектное решение, добавленное после аудита. {norm:sp8-5-2-tab2}\n",
            encoding="utf-8",
        )

        problems = audit_binding_issues(project)
        self.assertTrue(any("draft.md" in problem for problem in problems))
        self.assertEqual(Stage.DRAFT_READY, refresh_stage(project).stage)
        with self.assertRaises(RuntimeError):
            finalize_markdown(project)

        report = audit_project(project)
        self.assertTrue(report["passed"])
        self.assertEqual([], audit_binding_issues(project))
        self.assertEqual(Stage.AUDIT_PASSED, refresh_stage(project).stage)

    def test_audit_report_records_artifact_hashes(self) -> None:
        project = self.temp_project_parent() / "hash-binding"
        create_demo(project)

        report = io.read_json(io.artifact_paths(project).audit_report)
        self.assertEqual("mpb-pz-flow.audit_project", report["engine"])
        hashes = report["artifact_hashes"]
        self.assertEqual(io.sha256_path(io.artifact_paths(project).draft), hashes["draft.md"])
        self.assertEqual(io.sha256_path(io.artifact_paths(project).passport), hashes["passport.json"])
        self.assertIsNone(hashes["agent_findings.json"])

    def test_agent_findings_error_blocks_pass(self) -> None:
        project = self.temp_project_parent() / "agent-findings"
        create_demo(project)

        findings = {
            "level_2": [
                {
                    "code": "draft.declarative",
                    "message": "Абзац 2 декларативен: нет привязки к параметру объекта.",
                    "severity": "error",
                }
            ]
        }
        io.write_json(io.artifact_paths(project).agent_findings, findings)

        report = audit_project(project)
        self.assertFalse(report["passed"])
        self.assertEqual("ВЕРНУТЬ НА ПРАВКУ", report["verdict"])
        agent_issues = [issue for issue in report["level_2"] if issue["source"] == "agent"]
        self.assertEqual(1, len(agent_issues))
        self.assertEqual("draft.declarative", agent_issues[0]["code"])
        self.assertEqual(Stage.DRAFT_READY, refresh_stage(project).stage)

    def test_agent_findings_warning_does_not_block(self) -> None:
        project = self.temp_project_parent() / "agent-warning"
        create_demo(project)

        findings = {"level_2": [{"code": "style.note", "message": "Стилистическое замечание.", "severity": "warning"}]}
        io.write_json(io.artifact_paths(project).agent_findings, findings)

        report = audit_project(project)
        self.assertTrue(report["passed"])

    def test_malformed_agent_findings_block_pass(self) -> None:
        project = self.temp_project_parent() / "agent-broken"
        create_demo(project)

        io.artifact_paths(project).agent_findings.write_text("{broken json", encoding="utf-8")
        report = audit_project(project)
        self.assertFalse(report["passed"])
        codes = {issue["code"] for issue in report["level_1"]}
        self.assertIn("agent_findings.invalid", codes)

    def test_section_artifact_sets_are_isolated_and_validated(self) -> None:
        project = self.temp_project_parent() / "sections"
        create_demo(project)
        second_section = "Безопасность пожарных подразделений"
        slug = io.register_section(project, second_section)

        norm = NormEntry(
            norm_id="sp4-7-2",
            document="СП 4.13130",
            edition_year=2013,
            point="п. 7.2",
            quote="В зданиях высотой 10 и более метров должны предусматриваться выходы на кровлю.",
            subject="выходы на кровлю",
            trigger_parameter="height_m",
            source_file="standards/demo/SP_4_13130_2013_excerpt.md",
        )
        io.write_norms(project, [norm], section=second_section)
        io.write_matrix(
            project,
            [
                MatrixEntry(
                    norm_id="sp4-7-2",
                    document_point="СП 4.13130.2013, п. 7.2",
                    status="применимо",
                    passport_basis="высота здания 6.8 м",
                    numeric_thresholds="порог 10 м",
                    collisions="",
                    text_parameters="выходы на кровлю не требуются",
                )
            ],
            section=second_section,
        )

        section_paths = io.artifact_paths(project, second_section)
        self.assertIn(f"sections/{slug}", section_paths.norms.as_posix())
        self.assertNotEqual(io.artifact_paths(project).norms, section_paths.norms)

        section_paths.draft.write_text(
            "# Раздел\n\n[Тип А] Выходы на кровлю не требуются: высота 6.8 м менее 10 м. {norm:sp4-7-2}\n",
            encoding="utf-8",
        )
        issues = validate_draft(project, second_section)
        codes = {issue.code for issue in issues}
        self.assertIn("draft.visible_norm_reference", codes)
        self.assertTrue(all(f"sections/{slug}" in issue.artifact for issue in issues))

        report = audit_project(project, second_section)
        self.assertFalse(report["passed"])
        self.assertEqual(slug, report["section_slug"])
        # Аудит чужой секции не трогает стадию активной секции.
        self.assertEqual(Stage.AUDIT_PASSED, refresh_stage(project).stage)

    def test_compute_stage_per_section(self) -> None:
        project = self.temp_project_parent() / "stage-per-section"
        create_demo(project)
        second_section = "Категорирование"
        io.register_section(project, second_section)
        self.assertEqual(Stage.AUDIT_PASSED, compute_stage(project))
        self.assertEqual(Stage.PASSPORT_READY, compute_stage(project, second_section))

    def test_assemble_draft_for_named_section(self) -> None:
        project = self.temp_project_parent() / "assemble-section"
        create_demo(project)
        second_section = "Наружное ВПС и проезды (повтор)"
        io.register_section(project, second_section)
        io.write_norms(project, io.read_norms(project), section=second_section)
        io.write_matrix(project, io.read_matrix(project), section=second_section)

        draft = assemble_draft(project, second_section)
        self.assertIn(second_section, draft)
        self.assertTrue(io.artifact_paths(project, second_section).draft.exists())

    def test_gate_cli_passes_then_blocks_after_tamper(self) -> None:
        from mpb_pz_flow.cli import main as cli_main

        project = self.temp_project_parent() / "gate"
        create_demo(project)
        self.assertEqual(0, cli_main(["gate", "--project-dir", str(project)]))

        draft_path = io.artifact_paths(project).draft
        draft_path.write_text(draft_path.read_text(encoding="utf-8") + "\n", encoding="utf-8")
        self.assertEqual(2, cli_main(["gate", "--project-dir", str(project)]))

    def test_validate_strict_fails_on_warnings(self) -> None:
        from mpb_pz_flow.cli import main as cli_main

        project = self.temp_project_parent() / "strict"
        create_demo(project)
        # Убираем источники и манифест: пропавшие source_file дают warning.
        shutil.rmtree(project / "standards")
        self.assertEqual(0, cli_main(["validate", "--project-dir", str(project)]))
        self.assertEqual(2, cli_main(["validate", "--project-dir", str(project), "--strict"]))


class Phase1CorpusTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def _tamper_norm_quote(self, project: Path, norm_id: str, new_quote: str) -> None:
        rows = io.read_jsonl(io.artifact_paths(project).norms)
        for row in rows:
            if row["norm_id"] == norm_id:
                row["quote"] = new_quote
        io.write_jsonl(io.artifact_paths(project).norms, rows)

    def test_demo_quotes_are_verified_against_real_sources(self) -> None:
        project = self.temp_project_parent() / "verified-demo"
        create_demo(project)

        issues = validate_all(project)
        self.assertEqual([], [issue for issue in issues if issue.severity == "error"])
        codes = {issue.code for issue in issues}
        self.assertNotIn("norms.source_file_missing", codes)
        self.assertNotIn("norms.quote_unverified", codes)

    def test_fabricated_quote_blocks_pipeline(self) -> None:
        project = self.temp_project_parent() / "fake-quote"
        create_demo(project)

        self._tamper_norm_quote(
            project,
            "sp8-6-3",
            "Продолжительность тушения пожара следует принимать 6 часов для всех производственных зданий.",
        )

        issues = validate_quotes(project)
        codes = {issue.code for issue in issues}
        self.assertIn("norms.quote_unverified", codes)
        self.assertEqual(Stage.PASSPORT_READY, compute_stage(project))
        report = audit_project(project)
        self.assertFalse(report["passed"])

    def test_quote_with_ellipsis_matches_segments_in_order(self) -> None:
        source = corpus.normalize_for_match(
            "Число выходов на кровлю следует предусматривать по расчету. "
            "Дополнительные требования установлены для отдельных классов. "
            "По пожарным лестницам через каждые 200 метров по периметру зданий класса Ф5."
        )
        quote = "Число выходов на кровлю следует предусматривать по расчету ... через каждые 200 метров по периметру зданий класса Ф5."
        self.assertTrue(corpus.quote_matches(source, quote))
        reversed_quote = "через каждые 200 метров по периметру зданий класса Ф5 ... Число выходов на кровлю следует предусматривать по расчету."
        self.assertFalse(corpus.quote_matches(source, reversed_quote))

    def test_edition_not_accepted_is_error(self) -> None:
        project = self.temp_project_parent() / "edition-not-accepted"
        create_demo(project)

        rows = io.read_jsonl(io.artifact_paths(project).norms)
        for row in rows:
            if row["norm_id"] == "sp4-driveways":
                row["edition_year"] = 2016
        io.write_jsonl(io.artifact_paths(project).norms, rows)

        issues = validate_editions(project)
        codes = {issue.code for issue in issues}
        self.assertIn("norms.edition_not_accepted", codes)

    def test_two_editions_of_same_document_are_rejected(self) -> None:
        project = self.temp_project_parent() / "duplicate-edition"
        create_demo(project)

        decisions_path = io.artifact_paths(project).decisions
        decisions = io.read_json(decisions_path)
        decisions["standard_editions"].append({"document": "СП 8.13130", "edition_year": 2016})
        io.write_json(decisions_path, decisions)

        issues = validate_editions(project)
        codes = {issue.code for issue in issues}
        self.assertIn("decisions.duplicate_edition", codes)

    def test_collision_requires_resolution_record(self) -> None:
        project = self.temp_project_parent() / "collision"
        create_demo(project)

        rows = io.read_jsonl(io.artifact_paths(project).norms)
        for row in rows:
            if row["norm_id"] == "sp8-5-2-tab2":
                row["collision_with"] = "СП 31.13330"
        io.write_jsonl(io.artifact_paths(project).norms, rows)

        codes = {issue.code for issue in validate_editions(project)}
        self.assertIn("decisions.collision_unresolved", codes)

        decisions_path = io.artifact_paths(project).decisions
        decisions = io.read_json(decisions_path)
        decisions["collisions"].append(
            {
                "norms": ["sp8-5-2-tab2"],
                "description": "СП 8.13130 и СП 31.13330 нормируют расход воды по-разному.",
                "resolution": "Принят СП 8.13130.2020 как профильный документ системы противопожарной защиты.",
            }
        )
        io.write_json(decisions_path, decisions)

        codes = {issue.code for issue in validate_editions(project)}
        self.assertNotIn("decisions.collision_unresolved", codes)

    def test_ingest_cp1251_source_is_normalized_to_utf8(self) -> None:
        project = self.temp_project_parent() / "cp1251"
        create_demo(project)

        text = (
            "СП 9.13130.2009. Техника пожарная. Огнетушители. Требования к эксплуатации. "
            "Настоящий свод правил устанавливает требования к выбору, размещению, "
            "техническому обслуживанию и перезарядке переносных и передвижных огнетушителей."
        ) * 3
        raw = self.temp_project_parent() / "sp9_cp1251.txt"
        raw.write_bytes(text.encode("cp1251"))

        entry = corpus.ingest_file(project, raw, document="СП 9.13130", edition_year=2009)
        stored = corpus.standards_dir(project) / entry.file
        self.assertEqual(io.sha256_path(stored), entry.sha256)
        self.assertIn("огнетушителей", stored.read_text(encoding="utf-8"))
        self.assertIsNotNone(corpus.find_document(project, "СП 9.13130", 2009))

    def test_ingest_rejects_corrupted_source(self) -> None:
        project = self.temp_project_parent() / "corrupted"
        create_demo(project)

        mojibake = ("!  1.13130 >AA88 B@51>20=8O ?>60@=>9 157>?0A=>AB8 " * 20)
        raw = self.temp_project_parent() / "broken.md"
        raw.write_text(mojibake, encoding="utf-8")

        with self.assertRaises(corpus.CorpusError):
            corpus.ingest_file(project, raw, document="СП 1.13130", edition_year=2020)

    def test_corpus_tamper_is_detected_and_blocks_audit(self) -> None:
        project = self.temp_project_parent() / "corpus-tamper"
        create_demo(project)

        source = corpus.standards_dir(project) / "demo/SP_8_13130_2020_excerpt.md"
        source.write_text(source.read_text(encoding="utf-8").replace("3 часа", "2 часа"), encoding="utf-8")

        issues = corpus.verify_corpus(project)
        codes = {issue.code for issue in issues}
        self.assertIn("corpus.hash_mismatch", codes)
        report = audit_project(project)
        self.assertFalse(report["passed"])

    def test_manifest_change_invalidates_audit_binding(self) -> None:
        project = self.temp_project_parent() / "manifest-binding"
        create_demo(project)
        self.assertEqual([], audit_binding_issues(project))

        documents = corpus.read_manifest(project)
        documents[0].title = "Изменённое название"
        corpus.write_manifest(project, documents)

        problems = audit_binding_issues(project)
        self.assertTrue(any("standards/manifest.json" in problem for problem in problems))


class Phase2EngineeringTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    # --- Калькуляторы ---

    def test_f5_water_flow_hvo_case(self) -> None:
        from mpb_pz_flow.calculators import calc_external_water_flow_f5

        confirmed = {
            "building_volume_m3": 32678.8,
            "fire_resistance_degree": "III",
            "preliminary_fire_category": "В",
            "building_width_m": 36,
        }
        result = calc_external_water_flow_f5(confirmed)
        self.assertEqual("25", result.value)
        self.assertEqual("л/с", result.unit)
        self.assertIn("СП 8.13130.2020, табл. 3", result.basis)
        self.assertIn("32678,8 м3", result.formula.replace("32678.8", "32678,8"))

    def test_f5_water_flow_low_volume_category_d(self) -> None:
        from mpb_pz_flow.calculators import calc_external_water_flow_f5

        confirmed = {
            "building_volume_m3": 4000,
            "fire_resistance_degree": "II",
            "building_fire_category": "Д",
        }
        self.assertEqual("10", calc_external_water_flow_f5(confirmed).value)

    def test_f5_water_flow_refuses_outside_encoded_range(self) -> None:
        from mpb_pz_flow.calculators import CalcError, calc_external_water_flow_f5

        confirmed = {
            "building_volume_m3": 500000,
            "fire_resistance_degree": "III",
            "building_fire_category": "В",
        }
        with self.assertRaises(CalcError):
            calc_external_water_flow_f5(confirmed)

    def test_f5_water_flow_refuses_wide_building(self) -> None:
        from mpb_pz_flow.calculators import CalcError, calc_external_water_flow_f5

        confirmed = {
            "building_volume_m3": 30000,
            "fire_resistance_degree": "III",
            "building_fire_category": "В",
            "building_width_m": 72,
        }
        with self.assertRaises(CalcError):
            calc_external_water_flow_f5(confirmed)

    def test_truck_access_and_roof_ladders(self) -> None:
        from mpb_pz_flow.calculators import calc_fire_truck_access_sides, calc_roof_ladder_count

        confirmed = {"building_width_m": 36, "building_length_m": 78, "height_m": 12.18}
        sides = calc_fire_truck_access_sides(confirmed, "F5.1")
        self.assertIn("двух", sides.value)
        ladders = calc_roof_ladder_count(confirmed, "F5.1")
        self.assertEqual("2", ladders.value)
        self.assertIn("228", ladders.formula)

    def test_public_water_flow_demo_case(self) -> None:
        from mpb_pz_flow.calculators import calc_external_water_flow_public

        result = calc_external_water_flow_public({"building_volume_m3": 5400, "floors": 1})
        self.assertEqual("15", result.value)
        self.assertIn("табл. 2", result.basis)

    def test_drive_distance_three_tiers_and_clause(self) -> None:
        from mpb_pz_flow.calculators import calc_fire_drive_distance

        low = calc_fire_drive_distance({"height_m": 6.56})
        self.assertEqual("не более 25", low.value)  # ≤12 м — раньше баг давал «5–8»
        self.assertIn("п. 8.2.6", low.basis)        # раньше был неверный «п. 8.8»

        mid = calc_fire_drive_distance({"height_m": 20})
        self.assertEqual("5–8", mid.value)

        high = calc_fire_drive_distance({"height_m": 40})
        self.assertEqual("8–10", high.value)

    def test_fire_duration_clause_and_public_two_hours(self) -> None:
        from mpb_pz_flow.calculators import calc_fire_duration

        # производственное здание (Ф5) — 3 ч, корректный локатор п. 5.17 (был п. 6.3)
        prod = calc_fire_duration({"fire_resistance_degree": "IV", "structural_fire_hazard_class": "С0"}, "F5.1")
        self.assertEqual("3", prod.value)
        self.assertIn("п. 5.17", prod.basis)
        self.assertNotIn("6.3", prod.basis)

        # жилое/общественное I–II степени класса С0 — 2 ч
        public = calc_fire_duration({"fire_resistance_degree": "II", "structural_fire_hazard_class": "С0"}, "F1.3")
        self.assertEqual("2", public.value)

    def test_calc_run_writes_registry_and_demo_draft_has_table(self) -> None:
        project = self.temp_project_parent() / "calc-registry"
        create_demo(project)

        registry = io.read_json(io.artifact_paths(project).calculations)
        self.assertEqual("mpb-pz-flow.calculators", registry["engine"])
        ids = {result["calc_id"] for result in registry["results"]}
        self.assertIn("sp8_t2_external_flow_public", ids)
        self.assertNotIn("sp8_t3_external_flow_f5", ids)

        draft = io.artifact_paths(project).draft.read_text(encoding="utf-8")
        self.assertIn("## Расчетные показатели", draft)

    # --- Триггеры ---

    def test_trigger_mismatch_without_justification_is_error(self) -> None:
        from mpb_pz_flow.validators import validate_triggers

        project = self.temp_project_parent() / "trigger-mismatch"
        create_demo(project)

        matrix_path = io.artifact_paths(project).matrix
        rows = io.read_jsonl(matrix_path)
        for row in rows:
            if row["norm_id"] == "sp4-driveways":
                row["status"] = "неприменимо"
        io.write_jsonl(matrix_path, rows)

        codes = {issue.code for issue in validate_triggers(project)}
        self.assertIn("matrix.trigger_mismatch", codes)
        report = audit_project(project)
        self.assertFalse(report["passed"])

    def test_trigger_override_with_justification_is_warning(self) -> None:
        from mpb_pz_flow.validators import validate_triggers

        project = self.temp_project_parent() / "trigger-override"
        create_demo(project)

        matrix_path = io.artifact_paths(project).matrix
        rows = io.read_jsonl(matrix_path)
        for row in rows:
            if row["norm_id"] == "sp4-driveways":
                row["status"] = "требует инженерной проверки"
                row["passport_basis"] = "отсутствует схема генплана с проездами"
                row["override_justification"] = "Фактическая организация подъезда зависит от генплана, который не выпущен."
        io.write_jsonl(matrix_path, rows)

        issues = validate_triggers(project)
        self.assertTrue(issues)
        self.assertTrue(all(issue.severity == "warning" for issue in issues))
        self.assertEqual({"matrix.trigger_override"}, {issue.code for issue in issues})

    def test_build_initial_matrix_uses_engine_proposals(self) -> None:
        from mpb_pz_flow.pipeline import build_initial_matrix

        project = self.temp_project_parent() / "matrix-proposals"
        create_demo(project)

        rows = build_initial_matrix(project)
        by_id = {row.norm_id: row for row in rows}
        self.assertEqual("применимо", by_id["sp4-driveways"].status)
        self.assertIn("height_m", by_id["sp4-driveways"].passport_basis)
        # Нормы без структурных триггеров остаются на инженерной проверке.
        self.assertEqual("требует инженерной проверки", by_id["sp8-5-2-tab2"].status)

    def test_unknown_trigger_param_proposes_engineering_check(self) -> None:
        from mpb_pz_flow.triggers import propose_status

        norm = NormEntry(
            norm_id="x",
            document="СП 4.13130",
            edition_year=2013,
            point="п. 7.2",
            quote="x" * 30,
            subject="t",
            trigger_parameter="height_to_parapet_m",
            source_file="standards/none.md",
            triggers=[{"param": "height_to_parapet_m", "op": ">=", "value": 10}],
        )
        status, basis = propose_status(norm, {"confirmed": {}})
        self.assertEqual("требует инженерной проверки", status)
        self.assertIn("отсутств", basis.lower())

    # --- Числовая согласованность ---

    def test_unverified_unit_number_blocks(self) -> None:
        from mpb_pz_flow.validators import validate_numbers

        project = self.temp_project_parent() / "fake-number"
        create_demo(project)

        draft_path = io.artifact_paths(project).draft
        draft_path.write_text(
            draft_path.read_text(encoding="utf-8")
            + "\n[Тип Г] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 расход принят 97 л/с. {norm:sp8-5-2-tab2}\n",
            encoding="utf-8",
        )

        issues = validate_numbers(project)
        codes = {issue.code for issue in issues}
        self.assertIn("draft.number_unverified", codes)
        self.assertEqual(Stage.MATRIX_READY, compute_stage(project))

    def test_demo_numbers_all_have_provenance(self) -> None:
        from mpb_pz_flow.validators import validate_numbers

        project = self.temp_project_parent() / "clean-numbers"
        create_demo(project)
        self.assertEqual([], validate_numbers(project))

    def test_calculations_tamper_invalidates_audit(self) -> None:
        project = self.temp_project_parent() / "calc-tamper"
        create_demo(project)
        self.assertEqual([], audit_binding_issues(project))

        registry = io.read_json(io.artifact_paths(project).calculations)
        registry["results"][0]["value"] = "999"
        io.write_json(io.artifact_paths(project).calculations, registry)

        problems = audit_binding_issues(project)
        self.assertTrue(any("calculations.json" in problem for problem in problems))

    def test_finalize_converts_decimal_point_to_comma_before_units(self) -> None:
        project = self.temp_project_parent() / "decimal-comma"
        create_demo(project)

        final = finalize_markdown(project)
        self.assertIn("6,8 м", final)
        self.assertNotIn("6.8 м", final)
        # Номера пунктов не затрагиваются.
        self.assertIn("п. 5.2", final)


class Phase3TextQualityTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def _demo_with_draft(self, name: str, extra_paragraphs: str) -> Path:
        project = self.temp_project_parent() / name
        create_demo(project)
        draft_path = io.artifact_paths(project).draft
        draft_path.write_text(
            draft_path.read_text(encoding="utf-8") + "\n" + extra_paragraphs + "\n",
            encoding="utf-8",
        )
        return project

    def test_type_a_without_decision_verb_is_error(self) -> None:
        from mpb_pz_flow.style import validate_paragraph_types

        project = self._demo_with_draft(
            "grammar-a",
            "[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 объект хороший. {norm:sp8-5-2-tab2}",
        )
        codes = {issue.code for issue in validate_paragraph_types(project)}
        self.assertIn("draft.type_grammar", codes)

    def test_type_b_without_negation_is_error(self) -> None:
        from mpb_pz_flow.style import validate_paragraph_types

        project = self._demo_with_draft(
            "grammar-b",
            "[Тип Б] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 принято решение о подъезде. {norm:sp8-5-2-tab2}",
        )
        codes = {issue.code for issue in validate_paragraph_types(project)}
        self.assertIn("draft.type_grammar", codes)

    def test_type_g_without_unit_number_is_error(self) -> None:
        from mpb_pz_flow.style import validate_paragraph_types

        project = self._demo_with_draft(
            "grammar-g",
            "[Тип Г] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 расход принят по таблице. {norm:sp8-5-2-tab2}",
        )
        codes = {issue.code for issue in validate_paragraph_types(project)}
        self.assertIn("draft.type_grammar", codes)

    def test_demo_passes_type_grammar(self) -> None:
        from mpb_pz_flow.style import validate_paragraph_types, validate_style

        project = self.temp_project_parent() / "grammar-clean"
        create_demo(project)
        self.assertEqual([], validate_paragraph_types(project))
        self.assertEqual([], [issue for issue in validate_style(project) if issue.severity == "error"])

    def test_style_future_tense_is_error(self) -> None:
        from mpb_pz_flow.style import validate_style

        project = self._demo_with_draft(
            "style-future",
            "[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 будет предусмотрено пожаротушение. {norm:sp8-5-2-tab2}",
        )
        codes = {issue.code for issue in validate_style(project)}
        self.assertIn("draft.style_future", codes)

    def test_style_vague_citation_without_document_is_error(self) -> None:
        from mpb_pz_flow.style import validate_style

        project = self._demo_with_draft(
            "style-vague",
            "[Тип А] Согласно требованиям норм принято проектное решение об эвакуации. {norm:sp8-5-2-tab2}",
        )
        issues = [issue for issue in validate_style(project) if issue.code == "draft.style_vague_citation"]
        self.assertTrue(issues)
        self.assertEqual("error", issues[0].severity)

    def test_style_bare_sp_is_error(self) -> None:
        from mpb_pz_flow.style import validate_style

        project = self._demo_with_draft(
            "style-bare-sp",
            "[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 принято решение по требованиям данного СП. {norm:sp8-5-2-tab2}",
        )
        codes = {issue.code for issue in validate_style(project)}
        self.assertIn("draft.style_sp_without_number", codes)

    def test_style_document_without_edition_year_is_warning(self) -> None:
        from mpb_pz_flow.style import validate_style

        project = self._demo_with_draft(
            "style-no-year",
            "[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 принято решение; расход уточняется по СП 8.13130. {norm:sp8-5-2-tab2}",
        )
        issues = [issue for issue in validate_style(project) if issue.code == "draft.style_missing_edition_year"]
        self.assertTrue(issues)
        self.assertEqual("warning", issues[0].severity)

    def test_section_content_missing_topic_is_warning(self) -> None:
        from mpb_pz_flow.style import validate_section_content

        project = self.temp_project_parent() / "section-content"
        create_demo(project)
        draft_path = io.artifact_paths(project).draft
        # Оставляем только заголовок и один абзац без темы «подъезд».
        draft_path.write_text(
            "# Наружное ВПС и проезды\n\n"
            "[Тип Г] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 нормативный расход воды "
            "принят 15 л/с; продолжительность тушения 3 часа. {norm:sp8-5-2-tab2}\n",
            encoding="utf-8",
        )
        issues = validate_section_content(project)
        self.assertTrue(any("подъезд" in issue.message for issue in issues))
        self.assertTrue(all(issue.severity == "warning" for issue in issues))

    def test_demo_section_content_complete(self) -> None:
        from mpb_pz_flow.style import validate_section_content

        project = self.temp_project_parent() / "section-complete"
        create_demo(project)
        self.assertEqual([], validate_section_content(project))

    def test_stale_lacuna_in_draft_is_flagged(self) -> None:
        from mpb_pz_flow.style import validate_lacunae

        project = self._demo_with_draft(
            "lacuna-draft",
            "[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 принято решение; "
            "требования СП 4.13130 не раскрываются, поскольку источник по СП 4.13130 "
            "отсутствует в проектной базе. {norm:sp8-5-2-tab2}",
        )
        issues = [issue for issue in validate_lacunae(project) if issue.code == "draft.lacuna_stale"]
        self.assertTrue(issues)
        self.assertIn("СП 4.13130", issues[0].message)

    def test_stale_lacuna_in_decisions_is_flagged(self) -> None:
        from mpb_pz_flow.style import validate_lacunae

        project = self.temp_project_parent() / "lacuna-decisions"
        create_demo(project)
        decisions_path = io.artifact_paths(project).decisions
        decisions = io.read_json(decisions_path)
        decisions["lacunae"] = [
            {
                "parameter": "driveways_source",
                "status": "open",
                "note": "Локальный источник СП 4.13130 отсутствует, нормы по проездам не извлечены.",
            }
        ]
        io.write_json(decisions_path, decisions)

        issues = [issue for issue in validate_lacunae(project) if issue.code == "decisions.lacuna_stale"]
        self.assertTrue(issues)

    def test_resolved_lacuna_is_not_flagged(self) -> None:
        from mpb_pz_flow.style import validate_lacunae

        project = self.temp_project_parent() / "lacuna-resolved"
        create_demo(project)
        decisions_path = io.artifact_paths(project).decisions
        decisions = io.read_json(decisions_path)
        decisions["lacunae"] = [
            {"parameter": "driveways_source", "status": "resolved", "note": "СП 4.13130 добавлен в корпус."}
        ]
        io.write_json(decisions_path, decisions)
        self.assertEqual([], [issue for issue in validate_lacunae(project) if issue.code == "decisions.lacuna_stale"])

    def test_style_errors_block_stage_and_audit(self) -> None:
        project = self._demo_with_draft(
            "style-blocks",
            "[Тип А] В соответствии с СП 8.13130.2020, п. 5.2, табл. 2 будет предусмотрено решение. {norm:sp8-5-2-tab2}",
        )
        self.assertEqual(Stage.MATRIX_READY, compute_stage(project))
        report = audit_project(project)
        self.assertFalse(report["passed"])
        codes = {issue["code"] for issue in report["level_2"]}
        self.assertIn("draft.style_future", codes)


class Phase4AuditLoopTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def _break_quote(self, project: Path) -> None:
        rows = io.read_jsonl(io.artifact_paths(project).norms)
        for row in rows:
            if row["norm_id"] == "sp8-6-3":
                row["quote"] = "Выдуманная цитата, которой нет в источнике, длиной более двадцати символов."
        io.write_jsonl(io.artifact_paths(project).norms, rows)

    def _fix_quote(self, project: Path) -> None:
        rows = io.read_jsonl(io.artifact_paths(project).norms)
        for row in rows:
            if row["norm_id"] == "sp8-6-3":
                row["quote"] = (
                    "Продолжительность тушения пожара для расчета расхода воды на наружное "
                    "пожаротушение следует принимать 3 часа."
                )
        io.write_jsonl(io.artifact_paths(project).norms, rows)

    def test_routing_by_code_prefix(self) -> None:
        from mpb_pz_flow.routing import AGENT_1, AGENT_2, ORCHESTRATOR, route_code

        self.assertEqual(AGENT_1, route_code("norms.quote_unverified"))
        self.assertEqual(AGENT_1, route_code("passport.required"))
        self.assertEqual(AGENT_1, route_code("decisions.duplicate_edition"))
        self.assertEqual(AGENT_1, route_code("corpus.hash_mismatch"))
        self.assertEqual(AGENT_2, route_code("matrix.trigger_mismatch"))
        self.assertEqual(AGENT_2, route_code("draft.style_future"))
        self.assertEqual(ORCHESTRATOR, route_code("draft.number_unverified"))
        self.assertEqual(ORCHESTRATOR, route_code("agent_findings.invalid"))

    def test_report_issues_carry_route(self) -> None:
        project = self.temp_project_parent() / "routed"
        create_demo(project)
        self._break_quote(project)

        report = audit_project(project)
        self.assertFalse(report["passed"])
        quote_issues = [issue for issue in report["level_1"] if issue["code"] == "norms.quote_unverified"]
        self.assertTrue(quote_issues)
        self.assertEqual("agent_1", quote_issues[0]["route"])
        self.assertGreaterEqual(report["routing_summary"]["agent_1"], 1)

    def test_iteration_counts_only_changed_attempts(self) -> None:
        project = self.temp_project_parent() / "iterations"
        create_demo(project)
        self._break_quote(project)

        report = audit_project(project)
        self.assertEqual(1, report["iteration"])
        # Повторный прогон без правок — та же попытка.
        report = audit_project(project)
        self.assertEqual(1, report["iteration"])
        self.assertFalse(report["escalation_required"])

    def test_three_failed_iterations_escalate_with_persistent_findings(self) -> None:
        project = self.temp_project_parent() / "escalation"
        create_demo(project)
        self._break_quote(project)

        report = audit_project(project)
        self.assertEqual(1, report["iteration"])

        # Итерация 2: правка не устраняет дефект.
        passport_path = io.artifact_paths(project).passport
        passport = io.read_json(passport_path)
        passport["confirmed"]["note_attempt_2"] = "правка без устранения дефекта"
        io.write_json(passport_path, passport)
        report = audit_project(project)
        self.assertEqual(2, report["iteration"])
        self.assertFalse(report["escalation_required"])

        # Итерация 3: снова мимо.
        passport["confirmed"]["note_attempt_3"] = "ещё одна правка мимо"
        io.write_json(passport_path, passport)
        report = audit_project(project)
        self.assertEqual(3, report["iteration"])
        self.assertTrue(report["escalation_required"])
        self.assertEqual("ЭСКАЛАЦИЯ ПОЛЬЗОВАТЕЛЮ", report["verdict"])
        self.assertIn("norms.quote_unverified", report["persistent_findings"])
        self.assertEqual(3, len(report["iteration_history"]))

    def test_passing_audit_resets_loop(self) -> None:
        from mpb_pz_flow.pipeline import read_audit_loop

        project = self.temp_project_parent() / "loop-reset"
        create_demo(project)
        self._break_quote(project)
        audit_project(project)
        self.assertEqual(1, read_audit_loop(project)["iteration"])

        self._fix_quote(project)
        report = audit_project(project)
        self.assertTrue(report["passed"])
        loop = read_audit_loop(project)
        self.assertEqual(0, loop["iteration"])
        self.assertFalse(loop["escalated"])
        self.assertEqual(Stage.AUDIT_PASSED, refresh_stage(project).stage)


class Phase5ExportTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def test_abbreviations_collected_from_text(self) -> None:
        from mpb_pz_flow.front_matter import abbreviations_section, collect_abbreviations, unknown_abbreviations

        text = "Здание оборудуется СПС и СОУЭ; для ХВО предусмотрен ВПВ."
        found = dict(collect_abbreviations(text))
        self.assertIn("СПС", found)
        self.assertIn("СОУЭ", found)
        self.assertIn("ВПВ", found)
        self.assertIn("ХВО", unknown_abbreviations(text))
        section_md = abbreviations_section(text)
        self.assertIn("Перечень сокращений", section_md)
        self.assertIn("система пожарной сигнализации", section_md)

    def test_normative_documents_section_from_registry(self) -> None:
        from mpb_pz_flow.front_matter import normative_documents_section

        project = self.temp_project_parent() / "nd-list"
        create_demo(project)
        section_md = normative_documents_section(project)
        self.assertIn("СП 8.13130.2020", section_md)
        self.assertIn("добровольный", section_md)
        self.assertIn("Наружное противопожарное водоснабжение (демо-выписка)", section_md)
        self.assertIn("зарегистрирован", section_md)

    def test_write_front_matter_creates_file(self) -> None:
        from mpb_pz_flow.front_matter import write_front_matter

        project = self.temp_project_parent() / "fm-file"
        create_demo(project)
        target = write_front_matter(project)
        self.assertTrue(target.exists())
        self.assertIn("Перечень нормативно-правовых", target.read_text(encoding="utf-8"))

    def test_export_docx_gost_styling(self) -> None:
        try:
            from docx import Document
            from docx.shared import Mm
        except ImportError:
            self.skipTest("python-docx is not installed")
        from mpb_pz_flow.docx_export import export_docx

        project = self.temp_project_parent() / "gost-docx"
        create_demo(project)
        finalize_markdown(project)
        output = export_docx(project, title_page=True, front_matter=True)

        document = Document(str(output))
        section = document.sections[0]
        # Поля хранятся в twips — сравнение с допуском на округление.
        self.assertLess(abs(int(Mm(25)) - int(section.left_margin)), 1000)
        self.assertLess(abs(int(Mm(10)) - int(section.right_margin)), 1000)

        normal = document.styles["Normal"]
        self.assertEqual("Times New Roman", normal.font.name)

        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertIn("Демонстрационный магазин", texts)
        self.assertIn("МЕРОПРИЯТИЯ ПО ОБЕСПЕЧЕНИЮ ПОЖАРНОЙ БЕЗОПАСНОСТИ", texts)
        self.assertTrue(any(text.startswith("Таблица 1 —") for text in texts))
        self.assertTrue(any("Перечень нормативно-правовых" in text for text in texts))

        footer_xml = document.sections[0].footer.paragraphs[0]._p.xml
        self.assertIn("PAGE", footer_xml)

    def test_export_docx_without_title_page(self) -> None:
        try:
            from docx import Document
        except ImportError:
            self.skipTest("python-docx is not installed")
        from mpb_pz_flow.docx_export import export_docx

        project = self.temp_project_parent() / "no-title"
        create_demo(project)
        finalize_markdown(project)
        output = export_docx(project, title_page=False, front_matter=False)

        document = Document(str(output))
        texts = [paragraph.text for paragraph in document.paragraphs]
        self.assertNotIn("МЕРОПРИЯТИЯ ПО ОБЕСПЕЧЕНИЮ ПОЖАРНОЙ БЕЗОПАСНОСТИ", texts)


class Phase6WorkflowCommandsTests(unittest.TestCase):
    def temp_project_parent(self) -> Path:
        base = Path.cwd() / ".test-runs"
        base.mkdir(exist_ok=True)
        path = base / uuid.uuid4().hex
        path.mkdir()
        return path

    def test_norms_add_appends_then_updates_in_place(self) -> None:
        from mpb_pz_flow.cli import main as cli_main

        project = self.temp_project_parent() / "norms-add"
        create_demo(project)
        before = len(io.read_norms(project))

        new = project / "new_norm.jsonl"
        new.write_text(
            json.dumps(
                {
                    "norm_id": "test-extra-1",
                    "document": "СП 8.13130",
                    "edition_year": 2020,
                    "point": "п. 9.9",
                    "quote": "Тестовая норма достаточной длины для проверки команды norms-add.",
                    "subject": "тест",
                    "trigger_parameter": "building_volume_m3",
                    "source_file": "standards/demo/SP_8_13130_2020_excerpt.md",
                },
                ensure_ascii=False,
            ),
            encoding="utf-8",
        )

        self.assertEqual(0, cli_main(["norms-add", "--project-dir", str(project), "--file", str(new)]))
        self.assertEqual(before + 1, len(io.read_norms(project)))

        # повторная установка того же norm_id обновляет запись, а не плодит дубликат
        self.assertEqual(0, cli_main(["norms-add", "--project-dir", str(project), "--file", str(new)]))
        self.assertEqual(before + 1, len(io.read_norms(project)))

    def test_norms_add_rejects_missing_required_fields(self) -> None:
        from mpb_pz_flow.cli import main as cli_main

        project = self.temp_project_parent() / "norms-add-bad"
        create_demo(project)
        before = len(io.read_norms(project))

        bad = project / "bad.jsonl"
        bad.write_text(json.dumps({"norm_id": "x", "document": "СП 8.13130"}, ensure_ascii=False), encoding="utf-8")

        self.assertEqual(1, cli_main(["norms-add", "--project-dir", str(project), "--file", str(bad)]))
        self.assertEqual(before, len(io.read_norms(project)))  # ничего не записано

    def test_consolidate_includes_finalized_section_and_placeholders(self) -> None:
        from mpb_pz_flow.cli import main as cli_main
        from mpb_pz_flow.consolidate import build_consolidated_markdown

        project = self.temp_project_parent() / "consolidate"
        create_demo(project)
        finalize_markdown(project)  # даёт final.md активного раздела
        active = io.read_state(project).section

        self.assertEqual(0, cli_main(["consolidate", "--project-dir", str(project), "--no-docx"]))
        out = project / "artifacts" / "tom9_svod.md"
        self.assertTrue(out.exists())
        text = out.read_text(encoding="utf-8")
        self.assertIn("# Раздел 9", text)
        self.assertIn(io.read_state(project).object_name, text)

        md, included, placeholders = build_consolidated_markdown(project)
        self.assertIn("Общие положения", included)
        self.assertIn("Характеристика объекта", included)
        self.assertIn(active, included)  # финализированный содержательный раздел вставлен
        self.assertIn("Расчёт пожарных рисков", placeholders)
        self.assertNotIn("## Технико-экономические показатели", md)  # дублирующая ТЭП снята


if __name__ == "__main__":
    unittest.main()
