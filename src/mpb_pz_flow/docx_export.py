"""Экспорт ПЗ в DOCX с оформлением по ГОСТ Р 21.101 / ГОСТ 2.105.

Гарнитура Times New Roman, поля 25/10/20/20 мм, выравнивание по ширине
с абзацным отступом 12,5 мм, чёрные заголовки, подписи таблиц
«Таблица N — Наименование», титульный лист, нумерация страниц в нижнем
колонтитуле. Опционально — автогенерированная обвязка (перечень сокращений
и перечень НД) перед основным текстом.
"""

from __future__ import annotations

from pathlib import Path

from . import io
from .front_matter import build_front_matter
from .pipeline import audit_binding_issues, refresh_stage


def export_docx(
    project_dir: Path,
    section: str | None = None,
    title_page: bool = True,
    front_matter: bool = False,
) -> Path:
    try:
        from docx import Document
    except ImportError as exc:
        raise RuntimeError("python-docx is not installed. Install optional dependency: pip install .[docx]") from exc

    paths = io.artifact_paths(project_dir, section)
    problems = audit_binding_issues(project_dir, paths.section)
    if problems:
        raise RuntimeError("DOCX нельзя собрать: вердикт аудита недействителен. " + " ".join(problems))
    if not paths.final_md.exists():
        raise RuntimeError("final.md is missing. Run audit and finalize first.")

    text = paths.final_md.read_text(encoding="utf-8")
    if front_matter:
        matter = build_front_matter(project_dir, paths.section)
        if matter:
            text = matter + "\n" + text

    document = Document()
    _apply_gost_styles(document)
    _apply_page_setup(document)
    if title_page:
        state = io.read_state(project_dir)
        _add_title_page(document, state.object_name, paths.section)
    _render_markdown(document, text)
    _add_page_number_footer(document)

    document.save(paths.final_docx)
    refresh_stage(project_dir)
    return paths.final_docx


def _apply_gost_styles(document: object) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt, RGBColor

    normal = document.styles["Normal"]
    normal.font.name = "Times New Roman"
    normal.font.size = Pt(12)
    fmt = normal.paragraph_format
    fmt.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    fmt.first_line_indent = Mm(12.5)
    fmt.line_spacing = 1.15
    fmt.space_after = Pt(6)

    for style_name, size in (("Heading 1", 14), ("Heading 2", 13)):
        style = document.styles[style_name]
        style.font.name = "Times New Roman"
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        style.paragraph_format.first_line_indent = Mm(0)
        style.paragraph_format.space_before = Pt(12)
        style.paragraph_format.space_after = Pt(6)


def _apply_page_setup(document: object) -> None:
    from docx.shared import Mm

    for section in document.sections:
        section.left_margin = Mm(25)
        section.right_margin = Mm(10)
        section.top_margin = Mm(20)
        section.bottom_margin = Mm(20)


def _add_title_page(document: object, object_name: str, section_name: str) -> None:
    from datetime import date

    for _ in range(6):
        _centered(document, "", 12, False)
    _centered(document, object_name, 16, True)
    for _ in range(2):
        _centered(document, "", 12, False)
    _centered(document, "МЕРОПРИЯТИЯ ПО ОБЕСПЕЧЕНИЮ ПОЖАРНОЙ БЕЗОПАСНОСТИ", 14, True)
    _centered(document, "Раздел 9 проектной документации", 12, False)
    for _ in range(2):
        _centered(document, "", 12, False)
    _centered(document, section_name, 13, True)
    for _ in range(8):
        _centered(document, "", 12, False)
    _centered(document, str(date.today().year), 12, False)
    document.add_page_break()


def _centered(document: object, text: str, size: int, bold: bool) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm, Pt

    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Mm(0)
    run = paragraph.add_run(text)
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.font.bold = bold


def _render_markdown(document: object, text: str) -> None:
    lines = text.splitlines()
    table_counter = 0
    last_heading = ""
    index = 0
    while index < len(lines):
        line = lines[index]
        if line.startswith("# "):
            last_heading = line[2:].strip()
            document.add_heading(last_heading, level=1)
        elif line.startswith("## "):
            last_heading = line[3:].strip()
            document.add_heading(last_heading, level=2)
        elif line.startswith("|"):
            table_lines: list[str] = []
            while index < len(lines) and lines[index].startswith("|"):
                table_lines.append(lines[index])
                index += 1
            table_counter += 1
            _add_table_caption(document, table_counter, last_heading)
            _add_markdown_table(document, table_lines)
            continue
        elif line.strip():
            document.add_paragraph(line.strip())
        index += 1


def _add_table_caption(document: object, number: int, name: str) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    caption = f"Таблица {number} — {name}" if name else f"Таблица {number}"
    paragraph = document.add_paragraph(caption)
    paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
    paragraph.paragraph_format.first_line_indent = Mm(0)


def _add_page_number_footer(document: object) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    from docx.shared import Mm

    footer = document.sections[0].footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.first_line_indent = Mm(0)
    field = OxmlElement("w:fldSimple")
    field.set(qn("w:instr"), " PAGE ")
    paragraph._p.append(field)


def _add_markdown_table(document: object, lines: list[str]) -> None:
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.shared import Mm

    rows = [_parse_table_row(line) for line in lines if not _is_separator(line)]
    rows = [row for row in rows if row]
    if not rows:
        return

    table = document.add_table(rows=len(rows), cols=max(len(row) for row in rows))
    table.style = "Table Grid"
    for row_index, row in enumerate(rows):
        for col_index, value in enumerate(row):
            cell = table.cell(row_index, col_index)
            cell.text = value
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                paragraph.paragraph_format.first_line_indent = Mm(0)
                for run in paragraph.runs:
                    run.font.size = _pt(10)
                    if row_index == 0:
                        run.bold = True


def _parse_table_row(line: str) -> list[str]:
    return [cell.strip().strip("*") for cell in line.strip().strip("|").split("|")]


def _is_separator(line: str) -> bool:
    cells = _parse_table_row(line)
    return bool(cells) and all(set(cell.replace(":", "").replace("-", "").strip()) == set() for cell in cells)


def _pt(size: int) -> object:
    from docx.shared import Pt

    return Pt(size)
